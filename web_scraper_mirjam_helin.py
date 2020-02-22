# A program for scraping the contestant information on the Mirjam Helin website (2019). OBS! The site is changed and modified for every new competition.

# Importing modules (BeautifulSoup)
# In case of "ModuleNotFoundError: No module named 'exceptions'" --> pip install python_docx
from bs4 import BeautifulSoup, SoupStrainer
import requests
import urllib
from urllib.request import urlopen, Request
import docx # To install this, use pip install python_docx or pip install python-docx
from docx import Document
from docx import Document

# Importing packages with dashes between words
# import importlib  
# docx = importlib.import_module("python-docx")


# Choosing the url

url = "https://mirjamhelin.fi/fi/kilpailu/kilpailijat"

page = requests.get(url)
data = page.text
soup = BeautifulSoup(data)
linklist = []

# Finding links on the page, after that finding only the competitor page links and appending them to a list

for link in soup.find_all('a'):
    linklist.append(link.get('href'))
# print("The original link list is:", linklist)

temp = (linklist[107:199])
competitorlist = temp[::2]
# print(competitorlist)
competitorlinks = []

for i in competitorlist:
    competitorlinks.append("https://mirjamhelin.fi/" + i)
print("A list of competitor info page links:", competitorlinks, "\n")

# for i in competitorlinks:
## headers defines a dummy browser for this program to use. Without it an error occurs. 
headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/41.0.2228.0 Safari/537.3'}

for i in competitorlinks:
    reg_url = i
    # EXAMPLE: reg_url = "https://mirjamhelin.fi/fi/stefan-astakhov"
    req = Request(url=reg_url, headers=headers) 
    html = urlopen(req).read() 
    # print(html)  ### Prints also stuff not needed and prints umlauts wrong
    soup = BeautifulSoup(html)
    # html = urllib.request.urlopen("https://mirjamhelin.fi/fi/stefan-astakhov").read()

    # kill all script and style elements
    for script in soup(["script", "style"]):
        script.extract()    # rip it out

    # get text
    text = soup.get_text()

    # break into lines and remove leading and trailing space on each
    lines = (line.strip() for line in text.splitlines())
    # break multi-headlines into a line each
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    # drop blank lines
    text = '\n'.join(chunk for chunk in chunks if chunk)

    ## Editing the code to include only relevant text

    # print(text)
    print("\n")

    sub = "Takaisin"
    if sub in text:
        takaisinindex = (text.find(sub))
        newlineindex = (text.find("\n", takaisinindex))
        print(newlineindex)
        takaisin = text[takaisinindex:newlineindex]
    #     print(takaisin)


    sub2 = "CV"
    if sub2 in text:
        cvindex = (text.find(sub2))
        newlineindex2 = (text.find("\n", cvindex+3))
    #     print(newlineindex)
        title = text[cvindex+3:newlineindex2]
        print(title)


    sub = "Jaa"
    if sub in text:
        jaaindex = (text.find(sub))
        newlineindex3 = (text.find("\n", jaaindex-1))
        print("Newlineindex3 is:", newlineindex3)
        jaa = text[jaaindex:newlineindex3]
    #     print(jaa)

    text2 = text[newlineindex:newlineindex3]
#     print(text2)
    
    # Saving the files to your local folder

    document = Document()
    paragraph = document.add_paragraph(text2)
    document.save(title + ".docx") # replace ".docx" with ".txt" if you use the next block of code

    # Adding the texts to a word file

    # from docx import Document
    # document = Document()
    # paragraph = document.add_paragraph(text2)
    # document.save(title + ".docx")

# TO DO:
# - some of the files corrupted, others not
# - the amount of text scraped varies, thus, the file names are occasionally wrong


### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### ### 


### REGEX (does not seem to work)
# import re
# new_list = [x for x in linklist if re.search('fi', x)]
# for item in new_list:
#     print(item)

#### SUBSTRING SEARCH:
# sub = "a"
# for text in linklist:
#     if sub in text and sub is not None:
#         print(text)